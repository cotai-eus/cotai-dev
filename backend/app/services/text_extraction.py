"""
Serviço para extração de texto de diferentes tipos de documentos.
"""
import os
import tempfile
from typing import Optional, Dict, Any, List, Tuple
import pytesseract
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image

from app.core.config import settings

class TextExtractionService:
    """
    Serviço para extrair texto de diferentes tipos de documentos.
    """
    
    @staticmethod
    def extract_from_pdf(file_path: str, use_ocr: bool = False) -> Tuple[str, bool, int]:
        """
        Extrai texto de um arquivo PDF.
        
        Args:
            file_path: Caminho para o arquivo PDF
            use_ocr: Se deve usar OCR para extrair texto de imagens
            
        Returns:
            Tuple com (texto_extraído, usou_ocr, contagem_páginas)
        """
        pdf = PdfReader(file_path)
        text = ""
        used_ocr = False
        page_count = len(pdf.pages)
        
        # Primeiro tenta extrair texto diretamente
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n\n"
        
        # Se o texto extraído for muito curto e OCR está habilitado, tenta OCR
        if use_ocr and settings.OCR_ENABLED and len(text.strip()) < 100:
            temp_dir = tempfile.mkdtemp()
            try:
                # Converte páginas do PDF em imagens para OCR
                ocr_text = ""
                
                # Configura o caminho do Tesseract se especificado
                if settings.TESSERACT_PATH:
                    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
                
                # Usando PyMuPDF (fitz) para renderizar páginas do PDF (se disponível)
                try:
                    import fitz
                    doc = fitz.open(file_path)
                    for i, page in enumerate(doc):
                        pix = page.get_pixmap()
                        img_path = os.path.join(temp_dir, f"page_{i}.png")
                        pix.save(img_path)
                        
                        # Faz OCR na imagem
                        page_text = pytesseract.image_to_string(
                            Image.open(img_path),
                            lang=settings.OCR_LANGUAGE
                        )
                        ocr_text += page_text + "\n\n"
                    
                    # Usa texto OCR apenas se for mais longo que o texto extraído diretamente
                    if len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                        used_ocr = True
                except ImportError:
                    # Se PyMuPDF não estiver disponível, mantém o texto original
                    pass
            finally:
                # Limpa arquivos temporários
                import shutil
                shutil.rmtree(temp_dir)
                
        return text, used_ocr, page_count
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[str, int]:
        """
        Extrai texto de um arquivo DOCX.
        
        Args:
            file_path: Caminho para o arquivo DOCX
            
        Returns:
            Tuple com (texto_extraído, contagem_parágrafos)
        """
        doc = DocxDocument(file_path)
        text = ""
        
        # Extrai texto de parágrafos
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extrai texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text, len(doc.paragraphs)
    
    @staticmethod
    def extract_from_doc(file_path: str) -> Tuple[str, bool]:
        """
        Extrai texto de um arquivo DOC.
        Tenta usar antiword ou OCR como fallback.
        
        Args:
            file_path: Caminho para o arquivo DOC
            
        Returns:
            Tuple com (texto_extraído, usou_ocr)
        """
        text = ""
        used_ocr = False
        
        # Tenta usar 'antiword' se estiver disponível
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
            if result.returncode == 0:
                text = result.stdout
            else:
                # Se antiword falhar, tenta usar OCR
                if settings.OCR_ENABLED:
                    # Converte DOC para PDF e depois aplica OCR
                    try:
                        # Tenta usar LibreOffice para converter para PDF
                        temp_dir = tempfile.mkdtemp()
                        temp_pdf = os.path.join(temp_dir, "temp.pdf")
                        
                        convert_cmd = [
                            'libreoffice', '--headless', '--convert-to', 'pdf',
                            '--outdir', temp_dir, file_path
                        ]
                        
                        subprocess.run(convert_cmd, capture_output=True)
                        
                        if os.path.exists(temp_pdf):
                            # Extrai texto do PDF gerado
                            text, used_ocr, _ = TextExtractionService.extract_from_pdf(temp_pdf, use_ocr=True)
                    except Exception:
                        # Se a conversão falhar, não faz nada
                        pass
                    finally:
                        # Limpa arquivos temporários
                        import shutil
                        shutil.rmtree(temp_dir)
        except Exception:
            # Se antiword não estiver disponível ou falhar, tenta LibreOffice diretamente
            try:
                temp_dir = tempfile.mkdtemp()
                temp_txt = os.path.join(temp_dir, "temp.txt")
                
                convert_cmd = [
                    'libreoffice', '--headless', '--convert-to', 'txt:Text',
                    '--outdir', temp_dir, file_path
                ]
                
                subprocess.run(convert_cmd, capture_output=True)
                
                if os.path.exists(temp_txt):
                    with open(temp_txt, 'r', encoding='utf-8') as f:
                        text = f.read()
            except Exception:
                # Se tudo falhar, devolve texto vazio
                pass
            finally:
                # Limpa arquivos temporários
                import shutil
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
        
        return text, used_ocr
    
    @staticmethod
    def extract_from_excel(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Extrai dados de um arquivo Excel.
        
        Args:
            file_path: Caminho para o arquivo Excel
            
        Returns:
            Tuple com (texto_extraído, dados_planilha)
        """
        # Lê todas as planilhas
        excel_data = pd.read_excel(file_path, sheet_name=None)
        
        # Prepara saída de texto
        text = ""
        sheets_data = []
        
        # Processa cada planilha
        for sheet_name, df in excel_data.items():
            text += f"==== Planilha: {sheet_name} ====\n\n"
            
            # Converte cabeçalhos e dados para string
            headers = df.columns.tolist()
            text += ", ".join(str(h) for h in headers) + "\n"
            
            # Adiciona linhas
            for _, row in df.iterrows():
                text += ", ".join(str(v) for v in row.values) + "\n"
            
            text += "\n\n"
            
            # Armazena dados estruturados
            sheet_data = {
                "sheet_name": sheet_name,
                "headers": headers,
                "rows": df.to_dict('records')
            }
            sheets_data.append(sheet_data)
        
        return text, sheets_data
    
    @staticmethod
    def extract_from_csv(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Extrai dados de um arquivo CSV.
        
        Args:
            file_path: Caminho para o arquivo CSV
            
        Returns:
            Tuple com (texto_extraído, dados_csv)
        """
        # Tenta diferentes encodings e delimitadores comuns
        encodings = ['utf-8', 'latin1', 'cp1252']
        separators = [',', ';', '\t', '|']
        
        df = None
        for encoding in encodings:
            for sep in separators:
                try:
                    df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                    if len(df.columns) > 1:  # Se tem mais de uma coluna, encontramos o separador correto
                        break
                except Exception:
                    continue
            if df is not None and len(df.columns) > 1:
                break
        
        # Se todas as tentativas falharem, tenta uma leitura simples
        if df is None or len(df.columns) <= 1:
            try:
                df = pd.read_csv(file_path)
            except Exception:
                # Se tudo falhar, retorna texto vazio e lista vazia
                return "", []
        
        # Extrai texto
        text = ""
        
        # Adiciona cabeçalhos
        headers = df.columns.tolist()
        text += ", ".join(str(h) for h in headers) + "\n"
        
        # Adiciona linhas
        for _, row in df.iterrows():
            text += ", ".join(str(v) for v in row.values) + "\n"
        
        # Armazena dados estruturados
        csv_data = {
            "headers": headers,
            "rows": df.to_dict('records')
        }
        
        return text, [csv_data]
    
    @staticmethod
    def extract_from_image(file_path: str) -> str:
        """
        Extrai texto de uma imagem usando OCR.
        
        Args:
            file_path: Caminho para o arquivo de imagem
            
        Returns:
            Texto extraído
        """
        if not settings.OCR_ENABLED:
            return ""
        
        # Configura o caminho do Tesseract se especificado
        if settings.TESSERACT_PATH:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
        
        try:
            # Abre a imagem
            image = Image.open(file_path)
            
            # Pré-processamento para melhorar OCR
            # Converte para escala de cinza se não for
            if image.mode != 'L':
                image = image.convert('L')
            
            # Faz OCR
            text = pytesseract.image_to_string(image, lang=settings.OCR_LANGUAGE)
            return text
        except Exception as e:
            # Se houver erro no OCR, retorna texto vazio
            print(f"Erro no OCR: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text(file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Extrai texto de um arquivo, determinando o tipo de arquivo pelo MIME type.
        
        Args:
            file_path: Caminho para o arquivo
            mime_type: Tipo MIME do arquivo
            
        Returns:
            Dicionário com informações de extração
        """
        extraction_result = {
            "text": "",
            "used_ocr": False,
            "page_count": None,
            "structured_data": None,
            "error": None
        }
        
        try:
            # PDF
            if mime_type == 'application/pdf':
                extraction_result["text"], extraction_result["used_ocr"], extraction_result["page_count"] = \
                    TextExtractionService.extract_from_pdf(file_path, use_ocr=True)
            
            # DOCX
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                extraction_result["text"], extraction_result["page_count"] = \
                    TextExtractionService.extract_from_docx(file_path)
            
            # DOC
            elif mime_type == 'application/msword':
                extraction_result["text"], extraction_result["used_ocr"] = \
                    TextExtractionService.extract_from_doc(file_path)
            
            # XLSX
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                extraction_result["text"], extraction_result["structured_data"] = \
                    TextExtractionService.extract_from_excel(file_path)
            
            # XLS
            elif mime_type == 'application/vnd.ms-excel':
                extraction_result["text"], extraction_result["structured_data"] = \
                    TextExtractionService.extract_from_excel(file_path)
            
            # CSV
            elif mime_type == 'text/csv':
                extraction_result["text"], extraction_result["structured_data"] = \
                    TextExtractionService.extract_from_csv(file_path)
            
            # TXT
            elif mime_type == 'text/plain':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extraction_result["text"] = f.read()
            
            # Imagens
            elif mime_type in ['image/jpeg', 'image/png', 'image/tiff', 'image/bmp']:
                extraction_result["text"] = TextExtractionService.extract_from_image(file_path)
                extraction_result["used_ocr"] = True if extraction_result["text"] else False
            
            # Tipo não suportado
            else:
                extraction_result["error"] = f"Unsupported file type: {mime_type}"
                
        except Exception as e:
            extraction_result["error"] = str(e)
            
        return extraction_result
